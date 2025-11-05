#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文档处理图形界面应用（增强版）
功能：
1. 为指定目录结构下的最深一级子目录中的所有图片生成二维码
2. 将二维码插入到PDF中
3. 上传图片到阿里云OSS
"""

import os
import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext, simpledialog
import qrcode
from PIL import Image
from reportlab.lib.pagesizes import A3, A4, A5, landscape
from reportlab.pdfgen import canvas
from reportlab.lib.units import mm
import threading
import re
import shutil
from pypinyin import lazy_pinyin, Style
from oss_helper import OSSConfig, OSSUploader
from docx import Document
from docx.shared import Mm


def chinese_to_pinyin_initials(text):
    """
    将中文转换为拼音首字母，保留英文、数字和部分符号
    使用pypinyin库自动转换所有中文字符
    
    Args:
        text: 输入文本
        
    Returns:
        转换后的文本
    """
    result = []
    for char in text:
        if '\u4e00' <= char <= '\u9fff':  # 中文字符
            # 使用pypinyin获取首字母
            pinyin = lazy_pinyin(char, style=Style.FIRST_LETTER)
            if pinyin:
                result.append(pinyin[0].lower())
            else:
                result.append('x')  # 后备方案
        elif char.isalnum() or char in '-_.':
            # 保留英文、数字和常用符号
            result.append(char)
        elif char in ' ()（）':
            # 空格和括号转换为下划线
            result.append('_')
        # 其他字符忽略
    
    return ''.join(result)


def convert_path_to_pinyin(path):
    """
    将路径中的中文部分转换为拼音首字母
    
    Args:
        path: 文件路径
        
    Returns:
        转换后的路径
    """
    parts = path.split('/')
    converted_parts = [chinese_to_pinyin_initials(part) for part in parts]
    return '/'.join(converted_parts)


class OSSConfigDialog(tk.Toplevel):
    """OSS配置对话框"""
    
    def __init__(self, parent, config):
        super().__init__(parent)
        self.title("OSS配置")
        self.geometry("500x400")
        self.config = config
        self.result = False
        
        self.create_widgets()
        self.load_config()
        
        # 模态对话框
        self.transient(parent)
        self.grab_set()
        
    def create_widgets(self):
        """创建配置界面"""
        main_frame = ttk.Frame(self, padding="20")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # Access Key ID
        ttk.Label(main_frame, text="Access Key ID:").grid(row=0, column=0, sticky=tk.W, pady=5)
        self.access_key_id_var = tk.StringVar()
        ttk.Entry(main_frame, textvariable=self.access_key_id_var, width=50).grid(row=0, column=1, pady=5)
        
        # Access Key Secret
        ttk.Label(main_frame, text="Access Key Secret:").grid(row=1, column=0, sticky=tk.W, pady=5)
        self.access_key_secret_var = tk.StringVar()
        ttk.Entry(main_frame, textvariable=self.access_key_secret_var, width=50, show="-").grid(row=1, column=1, pady=5)
        
        # Endpoint
        ttk.Label(main_frame, text="Endpoint:").grid(row=2, column=0, sticky=tk.W, pady=5)
        self.endpoint_var = tk.StringVar()
        entry_endpoint = ttk.Entry(main_frame, textvariable=self.endpoint_var, width=50)
        entry_endpoint.grid(row=2, column=1, pady=5)
        # ttk.Label(main_frame, text="例: oss-cn-beijing.aliyuncs.com", 
                #  foreground="gray").grid(row=3, column=1, sticky=tk.W)
        
        # Bucket Name
        ttk.Label(main_frame, text="Bucket Name:").grid(row=4, column=0, sticky=tk.W, pady=5)
        self.bucket_name_var = tk.StringVar()
        ttk.Entry(main_frame, textvariable=self.bucket_name_var, width=50).grid(row=4, column=1, pady=5)
        
        # Base Path
        ttk.Label(main_frame, text="基础路径:").grid(row=5, column=0, sticky=tk.W, pady=5)
        self.base_path_var = tk.StringVar()
        entry_base_path = ttk.Entry(main_frame, textvariable=self.base_path_var, width=50)
        entry_base_path.grid(row=5, column=1, pady=5)
        # ttk.Label(main_frame, text="", 
                #  foreground="gray").grid(row=6, column=1, sticky=tk.W)
        
        # 按钮
        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=7, column=0, columnspan=2, pady=20)
        
        ttk.Button(button_frame, text="测试连接", command=self.test_connection).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="保存", command=self.save_config).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="取消", command=self.cancel).pack(side=tk.LEFT, padx=5)
        
#         # 说明
#         info_frame = ttk.LabelFrame(main_frame, text="说明", padding="10")
#         info_frame.grid(row=8, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=10)
        
#         info_text = """
# • Access Key可在阿里云控制台获取
# • Endpoint格式: oss-cn-<region>.aliyuncs.com
# • 配置将保存在本地文件中
# • 基础路径为OSS中的目录前缀，可以为空
#         """
#         ttk.Label(info_frame, text=info_text, justify=tk.LEFT).pack()
        
    def load_config(self):
        """加载配置"""
        self.access_key_id_var.set(self.config.access_key_id)
        self.access_key_secret_var.set(self.config.access_key_secret)
        self.endpoint_var.set(self.config.endpoint)
        self.bucket_name_var.set(self.config.bucket_name)
        self.base_path_var.set(self.config.base_path)
    
    def test_connection(self):
        """测试OSS连接"""
        # 临时保存配置
        self.save_to_config()
        
        # 测试连接
        uploader = OSSUploader(self.config)
        success, message = uploader.test_connection()
        
        if success:
            messagebox.showinfo("测试成功", "OSS连接测试成功！")
        else:
            messagebox.showerror("测试失败", f"OSS连接测试失败：\n{message}")
    
    def save_to_config(self):
        """保存到配置对象"""
        self.config.access_key_id = self.access_key_id_var.get().strip()
        self.config.access_key_secret = self.access_key_secret_var.get().strip()
        self.config.endpoint = self.endpoint_var.get().strip()
        self.config.bucket_name = self.bucket_name_var.get().strip()
        self.config.base_path = self.base_path_var.get().strip()
    
    def save_config(self):
        """保存配置"""
        self.save_to_config()
        
        if self.config.save_config():
            messagebox.showinfo("成功", "配置已保存！")
            self.result = True
            self.destroy()
        else:
            messagebox.showerror("错误", "保存配置失败！")
    
    def cancel(self):
        """取消"""
        self.destroy()


class DocumentProcessorApp:
    """文档处理应用主类"""
    
    def __init__(self, root):
        self.root = root
        self.root.title("文档处理工具V2.0")
        self.root.geometry("950x750")
        
        # 页面尺寸映射
        self.page_sizes = {
            "A3": A3,
            "A4": A4,
            "A5": A5,
            "自定义": None
        }
        
        # OSS配置
        self.oss_config = OSSConfig()
        self.oss_uploader = None
        if self.oss_config.is_valid():
            self.oss_uploader = OSSUploader(self.oss_config)
        
        # 创建界面
        self.create_widgets()
        
    def create_widgets(self):
        """创建GUI组件"""
        
        # 主框架
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # 第一行：目录设置和OSS设置
        first_row_container = ttk.Frame(main_frame)
        first_row_container.grid(row=0, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=5)
        first_row_container.columnconfigure(0, weight=1)
        first_row_container.columnconfigure(1, weight=1)
        
        # 目录选择
        dir_frame = ttk.LabelFrame(first_row_container, text="目录设置", padding="10")
        dir_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), padx=(0, 5))
        
        ttk.Label(dir_frame, text="根目录:").grid(row=0, column=0, sticky=tk.W, pady=5)
        self.root_dir_var = tk.StringVar()
        # 监听路径变化
        self.root_dir_var.trace_add('write', self.on_directory_changed)
        ttk.Entry(dir_frame, textvariable=self.root_dir_var, width=35).grid(row=0, column=1, padx=5, pady=5)
        ttk.Button(dir_frame, text="浏览...", command=self.browse_directory).grid(row=0, column=2, padx=5, pady=5)
        
        # OSS设置
        oss_frame = ttk.LabelFrame(first_row_container, text="OSS设置", padding="10")
        oss_frame.grid(row=0, column=1, sticky=(tk.W, tk.E, tk.N, tk.S), padx=(5, 0))
        
        oss_button_frame = ttk.Frame(oss_frame)
        oss_button_frame.grid(row=0, column=0, sticky=tk.W)
        
        ttk.Button(oss_button_frame, text="配置OSS", command=self.configure_oss).pack(side=tk.LEFT, padx=5)
        
        self.oss_status_var = tk.StringVar(value="未配置")
        self.update_oss_status()
        ttk.Label(oss_button_frame, textvariable=self.oss_status_var).pack(side=tk.LEFT, padx=10)
        
        # 上传选项
        self.auto_upload_var = tk.BooleanVar(value=True)
        ttk.Checkbutton(oss_frame, text="生成二维码后自动上传图片到OSS", 
                       variable=self.auto_upload_var).grid(row=1, column=0, sticky=tk.W, pady=5)
        
        # 第二行：目录结构和进度
        second_row_container = ttk.Frame(main_frame)
        second_row_container.grid(row=1, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=5)
        second_row_container.columnconfigure(0, weight=1)
        second_row_container.columnconfigure(1, weight=1)
        
        # 目录结构显示
        dir_type_frame = ttk.LabelFrame(second_row_container, text="目录结构", padding="10")
        dir_type_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), padx=(0, 5))
        
        ttk.Label(dir_type_frame, text="检测结果:").grid(row=0, column=0, sticky=tk.W, pady=5)
        self.dir_type_var = tk.StringVar(value="未检测")
        self.dir_type_label = ttk.Label(dir_type_frame, textvariable=self.dir_type_var, foreground="blue")
        self.dir_type_label.grid(row=0, column=1, sticky=tk.W, pady=5, padx=5)
        
        # 进度显示
        progress_frame = ttk.LabelFrame(second_row_container, text="进度", padding="10")
        progress_frame.grid(row=0, column=1, sticky=(tk.W, tk.E, tk.N, tk.S), padx=(5, 0))
        
        self.progress_var = tk.StringVar(value="就绪")
        ttk.Label(progress_frame, textvariable=self.progress_var).pack(anchor=tk.W)
        
        self.progress_bar = ttk.Progressbar(progress_frame, mode='indeterminate')
        self.progress_bar.pack(fill=tk.X, pady=5)
        
        # 文档类型选择
        doc_type_frame = ttk.LabelFrame(main_frame, text="文档类型", padding="10")
        doc_type_frame.grid(row=2, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=5)
        
        ttk.Label(doc_type_frame, text="输出格式:").grid(row=0, column=0, sticky=tk.W, pady=5)
        self.doc_type_var = tk.StringVar(value="PDF")
        doc_type_option_frame = ttk.Frame(doc_type_frame)
        doc_type_option_frame.grid(row=0, column=1, sticky=tk.W, pady=5)
        ttk.Radiobutton(doc_type_option_frame, text="PDF", variable=self.doc_type_var, 
                       value="PDF", command=self.on_doc_type_change).pack(side=tk.LEFT, padx=5)
        ttk.Radiobutton(doc_type_option_frame, text="Word", variable=self.doc_type_var, 
                       value="Word", command=self.on_doc_type_change).pack(side=tk.LEFT, padx=5)
        
        # PDF设置
        pdf_frame = ttk.LabelFrame(main_frame, text="PDF设置", padding="10")
        pdf_frame.grid(row=3, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=5)
        
        # 页面尺寸
        ttk.Label(pdf_frame, text="页面尺寸:").grid(row=0, column=0, sticky=tk.W, pady=5)
        self.page_size_var = tk.StringVar(value="自定义")
        page_size_combo = ttk.Combobox(pdf_frame, textvariable=self.page_size_var, 
                                       values=list(self.page_sizes.keys()), state="readonly", width=15)
        page_size_combo.grid(row=0, column=1, sticky=tk.W, padx=5, pady=5)
        page_size_combo.bind("<<ComboboxSelected>>", self.on_page_size_change)
        
        # 页面方向
        ttk.Label(pdf_frame, text="页面方向:").grid(row=1, column=0, sticky=tk.W, pady=5)
        self.page_orientation_var = tk.StringVar(value="纵向")
        orientation_frame = ttk.Frame(pdf_frame)
        orientation_frame.grid(row=1, column=1, sticky=tk.W, pady=5)
        ttk.Radiobutton(orientation_frame, text="纵向", variable=self.page_orientation_var, value="纵向").pack(side=tk.LEFT, padx=5)
        ttk.Radiobutton(orientation_frame, text="横向", variable=self.page_orientation_var, value="横向").pack(side=tk.LEFT, padx=5)
        
        # 自定义尺寸
        self.custom_size_frame = ttk.Frame(pdf_frame)
        self.custom_size_frame.grid(row=0, column=2, columnspan=4, sticky=tk.W, padx=5, pady=5)
        
        ttk.Label(self.custom_size_frame, text="宽度(mm):").pack(side=tk.LEFT, padx=2)
        self.custom_width_var = tk.StringVar(value="325")
        ttk.Entry(self.custom_size_frame, textvariable=self.custom_width_var, width=8).pack(side=tk.LEFT, padx=2)
        
        ttk.Label(self.custom_size_frame, text="高度(mm):").pack(side=tk.LEFT, padx=2)
        self.custom_height_var = tk.StringVar(value="238")
        ttk.Entry(self.custom_size_frame, textvariable=self.custom_height_var, width=8).pack(side=tk.LEFT, padx=2)
        
        # Word设置
        word_frame = ttk.LabelFrame(main_frame, text="Word设置", padding="10")
        word_frame.grid(row=4, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=5)
        
        # Word页面尺寸
        ttk.Label(word_frame, text="页面尺寸:").grid(row=0, column=0, sticky=tk.W, pady=5)
        self.word_size_var = tk.StringVar(value="自定义")
        word_size_combo = ttk.Combobox(word_frame, textvariable=self.word_size_var, 
                                       values=["A3", "A4", "A5", "自定义"], state="readonly", width=15)
        word_size_combo.grid(row=0, column=1, sticky=tk.W, padx=5, pady=5)
        word_size_combo.bind("<<ComboboxSelected>>", self.on_word_size_change)
        
        # Word自定义尺寸
        self.word_custom_size_frame = ttk.Frame(word_frame)
        self.word_custom_size_frame.grid(row=0, column=2, columnspan=4, sticky=tk.W, padx=5, pady=5)
        
        ttk.Label(self.word_custom_size_frame, text="宽度(mm):").pack(side=tk.LEFT, padx=2)
        self.word_width_var = tk.StringVar(value="340")
        ttk.Entry(self.word_custom_size_frame, textvariable=self.word_width_var, width=8).pack(side=tk.LEFT, padx=2)
        
        ttk.Label(self.word_custom_size_frame, text="高度(mm):").pack(side=tk.LEFT, padx=2)
        self.word_height_var = tk.StringVar(value="240")
        ttk.Entry(self.word_custom_size_frame, textvariable=self.word_height_var, width=8).pack(side=tk.LEFT, padx=2)
        
        # 默认隐藏Word设置
        word_frame.grid_remove()
        
        # 保存frame引用以便后续切换显示
        self.pdf_frame = pdf_frame
        self.word_frame = word_frame
        
        # 二维码设置
        qr_frame = ttk.LabelFrame(main_frame, text="二维码设置", padding="10")
        qr_frame.grid(row=5, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=5)
        
        # 二维码大小
        ttk.Label(qr_frame, text="二维码大小(mm):").grid(row=0, column=0, sticky=tk.W, pady=5)
        self.qr_size_var = tk.StringVar(value="20")
        ttk.Entry(qr_frame, textvariable=self.qr_size_var, width=10).grid(row=0, column=1, sticky=tk.W, padx=5, pady=5)
        
        # 二维码位置
        ttk.Label(qr_frame, text="X坐标(mm):").grid(row=1, column=0, sticky=tk.W, pady=5)
        self.qr_x_var = tk.StringVar(value="10")
        ttk.Entry(qr_frame, textvariable=self.qr_x_var, width=10).grid(row=1, column=1, sticky=tk.W, padx=5, pady=5)
        
        ttk.Label(qr_frame, text="Y坐标(mm):").grid(row=1, column=2, sticky=tk.W, pady=5, padx=(20, 0))
        self.qr_y_var = tk.StringVar(value="10")
        ttk.Entry(qr_frame, textvariable=self.qr_y_var, width=10).grid(row=1, column=3, sticky=tk.W, padx=5, pady=5)
        
        ttk.Label(qr_frame, text="注：坐标为二维码左下角位置，从页面左下角开始计算", 
                 foreground="blue").grid(row=2, column=0, columnspan=4, sticky=tk.W, pady=5)
        
        # 操作按钮
        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=6, column=0, columnspan=2, pady=15)
        
        self.start_button = ttk.Button(button_frame, text="开始处理", command=self.start_processing, width=15)
        self.start_button.pack(side=tk.LEFT, padx=5)
        
        ttk.Button(button_frame, text="清除日志", command=self.clear_log, width=15).pack(side=tk.LEFT, padx=5)
        
        # 日志显示
        log_frame = ttk.LabelFrame(main_frame, text="处理日志", padding="10")
        log_frame.grid(row=7, column=0, columnspan=2, sticky=(tk.W, tk.E, tk.N, tk.S), pady=5)
        
        self.log_text = scrolledtext.ScrolledText(log_frame, height=15, width=100)
        self.log_text.pack(fill=tk.BOTH, expand=True)
        
        # 配置网格权重
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        main_frame.columnconfigure(1, weight=1)
        main_frame.rowconfigure(7, weight=1)
        
    def update_oss_status(self):
        """更新OSS状态显示"""
        if self.oss_config.is_valid():
            self.oss_status_var.set(f"✓ 已配置 (Bucket: {self.oss_config.bucket_name})")
        else:
            self.oss_status_var.set("✗ 未配置")
    
    def configure_oss(self):
        """配置OSS"""
        dialog = OSSConfigDialog(self.root, self.oss_config)
        self.root.wait_window(dialog)
        
        # 更新上传器
        if self.oss_config.is_valid():
            self.oss_uploader = OSSUploader(self.oss_config)
        
        self.update_oss_status()
    
    def on_page_size_change(self, event=None):
        """页面尺寸改变时的回调"""
        if self.page_size_var.get() == "自定义":
            self.custom_size_frame.grid()
        else:
            self.custom_size_frame.grid_remove()
    
    def on_word_size_change(self, event=None):
        """Word页面尺寸改变时的回调"""
        if self.word_size_var.get() == "自定义":
            self.word_custom_size_frame.grid()
        else:
            self.word_custom_size_frame.grid_remove()
    
    def on_doc_type_change(self):
        """文档类型改变时的回调"""
        if self.doc_type_var.get() == "PDF":
            self.pdf_frame.grid()
            self.word_frame.grid_remove()
        else:  # Word
            self.pdf_frame.grid_remove()
            self.word_frame.grid()
    
    def on_directory_changed(self, *args):
        """当目录路径改变时的回调"""
        directory = self.root_dir_var.get().strip()
        if directory and os.path.isdir(directory):
            # 自动检测目录结构
            detected_type = self.detect_directory_type(directory)
            self.dir_type_var.set(detected_type)
            self.log(f"检测到目录: {directory}")
            self.log(f"自动检测目录结构: {detected_type}")
        elif directory:
            # 路径不存在
            self.dir_type_var.set("路径无效")
        else:
            # 路径为空
            self.dir_type_var.set("未检测")
    
    def browse_directory(self):
        """浏览并选择目录"""
        directory = filedialog.askdirectory(title="选择根目录")
        if directory:
            # trace_add已经会触发on_directory_changed，所以这里只需要设置值
            self.root_dir_var.set(directory)
    
    def log(self, message):
        """添加日志信息"""
        self.log_text.insert(tk.END, f"{message}\n")
        self.log_text.see(tk.END)
        self.root.update_idletasks()
    
    def clear_log(self):
        """清除日志"""
        self.log_text.delete(1.0, tk.END)
    
    def detect_directory_type(self, root_dir):
        """
        自动检测目录结构类型
        
        Args:
            root_dir: 根目录路径
            
        Returns:
            "村" 或 "乡"
        """
        try:
            # 检查是否有图片文件的函数
            def has_images(directory):
                image_extensions = {'.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp'}
                for item in os.listdir(directory):
                    if os.path.isfile(os.path.join(directory, item)):
                        _, ext = os.path.splitext(item)
                        if ext.lower() in image_extensions:
                            return True
                return False
            
            # 检查第一级子目录
            level1_dirs = [d for d in os.listdir(root_dir) 
                          if os.path.isdir(os.path.join(root_dir, d)) and not d.startswith('.')]
            
            if not level1_dirs:
                return "村"  # 默认返回村
            
            # 检查第一个一级子目录
            first_level1_path = os.path.join(root_dir, level1_dirs[0])
            
            # 如果一级目录中有图片，说明是二级结构（村）
            if has_images(first_level1_path):
                return "村"
            
            # 检查一级目录下是否有子目录
            level2_dirs = [d for d in os.listdir(first_level1_path)
                          if os.path.isdir(os.path.join(first_level1_path, d)) and not d.startswith('.')]
            
            if not level2_dirs:
                return "村"  # 一级目录下没有子目录，说明是村
            
            # 检查第一个二级目录
            first_level2_path = os.path.join(first_level1_path, level2_dirs[0])
            
            # 如果二级目录中有图片，说明是三级结构（乡）
            if has_images(first_level2_path):
                return "乡"
            
            # 默认返回村
            return "村"
            
        except Exception as e:
            self.log(f"检测目录结构失败: {str(e)}")
            return "村"  # 出错时默认返回村
    
    def get_page_size(self):
        """获取页面尺寸"""
        page_size_name = self.page_size_var.get()
        orientation = self.page_orientation_var.get()
        
        if page_size_name == "自定义":
            try:
                width = float(self.custom_width_var.get()) * mm
                height = float(self.custom_height_var.get()) * mm
                # 应用方向
                if orientation == "横向":
                    width, height = height, width
                return (width, height)
            except ValueError:
                messagebox.showerror("错误", "自定义尺寸必须是数字")
                return None
        else:
            page_size = self.page_sizes[page_size_name]
            # 应用方向
            if orientation == "横向":
                page_size = landscape(page_size)
            return page_size
    
    def get_word_page_size(self):
        """获取Word页面尺寸（毫米）"""
        word_size_name = self.word_size_var.get()
        
        # 标准尺寸（毫米）
        word_sizes = {
            "A3": (420, 297),
            "A4": (297, 210),
            "A5": (210, 148)
        }
        
        if word_size_name == "自定义":
            try:
                width = float(self.word_width_var.get())
                height = float(self.word_height_var.get())
                return (width, height)
            except ValueError:
                messagebox.showerror("错误", "Word自定义尺寸必须是数字")
                return None, None
        else:
            return word_sizes[word_size_name]
    
    def get_target_directories(self, root_dir, dir_type):
        """
        获取目标目录列表（最深一级子目录）
        
        Args:
            root_dir: 根目录路径
            dir_type: 目录类型，"村"（二级）或"乡"（三级）
            
        Returns:
            目标目录列表
        """
        target_dirs = []
        
        if dir_type == "村":
            # 二级目录结构：根目录/一级目录
            for item in os.listdir(root_dir):
                item_path = os.path.join(root_dir, item)
                if os.path.isdir(item_path) and not item.startswith('.'):
                    target_dirs.append(item_path)
        else:  # 乡
            # 三级目录结构：根目录/一级目录/二级目录
            for level1 in os.listdir(root_dir):
                level1_path = os.path.join(root_dir, level1)
                if os.path.isdir(level1_path) and not level1.startswith('.'):
                    for level2 in os.listdir(level1_path):
                        level2_path = os.path.join(level1_path, level2)
                        if os.path.isdir(level2_path) and not level2.startswith('.'):
                            target_dirs.append(level2_path)
        
        return target_dirs
    
    def get_images_in_directory(self, directory):
        """
        获取目录中的所有图片文件
        
        Args:
            directory: 目录路径
            
        Returns:
            图片文件完整路径列表
        """
        image_extensions = {'.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp'}
        images = []
        
        try:
            for item in os.listdir(directory):
                full_path = os.path.join(directory, item)
                if os.path.isfile(full_path):
                    _, ext = os.path.splitext(item)
                    if ext.lower() in image_extensions:
                        images.append(full_path)  # 返回完整路径而不是文件名
        except Exception as e:
            self.log(f"读取目录 {directory} 时出错: {str(e)}")
        
        return images
    
    def merge_images_horizontally(self, image_paths, output_path):
        """
        将多张图片垂直拼接成一张大图（上下排列），保持原图方向不变
        
        Args:
            image_paths: 图片路径列表
            output_path: 输出文件路径
        
        Returns:
            bool: 是否成功
        """
        try:
            if not image_paths:
                return False
            
            # 打开所有图片并根据EXIF信息自动旋转到正确方向
            images = []
            for img_path in image_paths:
                img = Image.open(img_path)
                # 处理EXIF方向信息，确保图片按正确方向显示
                try:
                    from PIL import ImageOps
                    img = ImageOps.exif_transpose(img)
                except Exception:
                    pass  # 如果没有EXIF信息或处理失败，保持原样
                images.append(img)
            
            # 计算拼接后的尺寸
            # 宽度取所有图片中的最大宽度
            max_width = max(img.width for img in images)
            # 高度为所有图片高度之和
            total_height = sum(img.height for img in images)
            
            # 创建新图片（使用RGB模式以支持JPEG保存）
            merged = Image.new('RGB', (max_width, total_height), (255, 255, 255))
            
            # 逐个粘贴图片（从上到下）
            y_offset = 0
            for img in images:
                # 如果图片有透明通道，转换为RGB
                if img.mode in ('RGBA', 'LA', 'P'):
                    # 创建白色背景
                    background = Image.new('RGB', img.size, (255, 255, 255))
                    if img.mode == 'P':
                        img = img.convert('RGBA')
                    background.paste(img, mask=img.split()[-1] if img.mode in ('RGBA', 'LA') else None)
                    img = background
                elif img.mode != 'RGB':
                    img = img.convert('RGB')
                
                # 将图片居中粘贴（如果宽度不一致）
                x_offset = (max_width - img.width) // 2
                merged.paste(img, (x_offset, y_offset))
                y_offset += img.height
            
            # 拼接完成（保持原图方向，按从上到下顺序，不做自动旋转）
            self.log(f"拼接完成，尺寸: {merged.width}x{merged.height}")
            
            # 保存为高质量JPEG（无损PNG太大）
            merged.save(output_path, 'JPEG', quality=95, optimize=True)
            
            # 关闭所有图片
            for img in images:
                img.close()
            
            return True
        except Exception as e:
            self.log(f"合并图片失败: {str(e)}")
            return False

    def generate_qrcode(self, url, output_path, size_mm=50):
        """
        生成二维码图片（优化后更小）
        
        Args:
            url: 二维码内容（URL）
            output_path: 输出文件路径
            size_mm: 二维码大小（毫米）
        """
        try:
            # 将毫米转换为像素（假设300 DPI）
            dpi = 300
            size_px = int(size_mm * dpi / 25.4)

            # 先以最小box_size生成二维码矩阵，再根据目标像素计算最优box_size
            border = 1  # 最小安全边框
            qr = qrcode.QRCode(
                version=1,
                error_correction=qrcode.constants.ERROR_CORRECT_L,
                box_size=1,
                border=border,
            )
            qr.add_data(url)
            qr.make(fit=True)

            # 获取模块数量（每边的模块数）
            modules = getattr(qr, 'modules_count', None)
            if modules is None:
                try:
                    modules = len(qr.get_matrix())
                except Exception:
                    modules = 21

            # 计算合适的box_size，使得生成时模块尽可能小
            # 保证至少为1像素
            box_size = max(1, size_px // (modules + 2 * border))

            # 使用计算出的box_size生成图像
            qr.box_size = box_size
            img = qr.make_image(fill_color="black", back_color="white")

            # 如果生成的尺寸与目标尺寸不同，使用最近邻插值放缩以保持模块清晰
            final_w = img.size[0]
            if final_w != size_px:
                img = img.resize((size_px, size_px), resample=Image.Resampling.NEAREST)

            img.save(output_path)
            return True
        except Exception as e:
            self.log(f"生成二维码失败: {str(e)}")
            return False
    
    def create_pdf_with_qrcode(self, qr_image_path, pdf_path, page_size, qr_size_mm, x_mm, y_mm):
        """
        创建PDF并插入二维码
        
        Args:
            qr_image_path: 二维码图片路径
            pdf_path: PDF输出路径
            page_size: 页面尺寸
            qr_size_mm: 二维码大小（毫米）
            x_mm: X坐标（毫米，二维码左下角）
            y_mm: Y坐标（毫米，二维码左下角）
        """
        try:
            c = canvas.Canvas(pdf_path, pagesize=page_size)
            
            # 将毫米转换为点（ReportLab使用点作为单位）
            qr_size = qr_size_mm * mm
            x_pos = x_mm * mm
            y_pos = y_mm * mm
            
            # 在PDF上绘制二维码
            c.drawImage(qr_image_path, x_pos, y_pos, width=qr_size, height=qr_size)
            
            c.save()
            return True
        except Exception as e:
            self.log(f"创建PDF失败: {str(e)}")
            return False
    
    def create_word_with_qrcode(self, qr_image_path, word_path, page_width_mm, page_height_mm, qr_size_mm, x_mm, y_mm):
        """
        创建Word并插入二维码
        
        Args:
            qr_image_path: 二维码图片路径
            word_path: Word输出路径
            page_width_mm: 页面宽度（毫米）
            page_height_mm: 页面高度（毫米）
            qr_size_mm: 二维码大小（毫米）
            x_mm: X坐标（毫米，二维码左上角）
            y_mm: Y坐标（毫米，二维码左上角）
        """
        try:
            doc = Document()
            
            # 设置页面尺寸
            section = doc.sections[0]
            section.page_width = Mm(page_width_mm)
            section.page_height = Mm(page_height_mm)
            
            # 设置最小边距
            section.top_margin = Mm(0)
            section.bottom_margin = Mm(0)
            section.left_margin = Mm(0)
            section.right_margin = Mm(0)
            
            # 添加段落（用于定位）
            paragraph = doc.add_paragraph()
            
            # 设置段落格式
            paragraph.paragraph_format.space_before = Mm(y_mm)
            paragraph.paragraph_format.left_indent = Mm(x_mm)
            
            # 插入二维码图片
            run = paragraph.add_run()
            run.add_picture(qr_image_path, width=Mm(qr_size_mm))
            
            # 保存Word文档
            doc.save(word_path)
            return True
        except Exception as e:
            self.log(f"创建Word失败: {str(e)}")
            return False
    
    def generate_index_html(self, directory, uploaded_files, dir_name):
        """
        生成索引HTML文件，用于在浏览器中查看图片列表
        
        Args:
            directory: 本地目录路径
            uploaded_files: 已上传的文件列表
            dir_name: 目录名称
            
        Returns:
            index.html文件路径
        """
        html_content = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{dir_name} - 图片浏览</title>
    <style>
        * {{
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }}
        body {{
            font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, "Helvetica Neue", Arial, sans-serif;
            background: #f5f5f5;
            padding: 20px;
        }}
        .container {{
            max-width: 1200px;
            margin: 0 auto;
            background: white;
            border-radius: 8px;
            box-shadow: 0 2px 8px rgba(0,0,0,0.1);
            padding: 30px;
        }}
        h1 {{
            color: #333;
            margin-bottom: 10px;
            font-size: 28px;
        }}
        .info {{
            color: #666;
            margin-bottom: 30px;
            font-size: 14px;
        }}
        .gallery {{
            display: grid;
            grid-template-columns: repeat(auto-fill, minmax(250px, 1fr));
            gap: 20px;
            margin-top: 20px;
        }}
        .image-item {{
            background: #fff;
            border: 1px solid #e0e0e0;
            border-radius: 8px;
            overflow: hidden;
            transition: transform 0.2s, box-shadow 0.2s;
            cursor: pointer;
        }}
        .image-item:hover {{
            transform: translateY(-4px);
            box-shadow: 0 4px 12px rgba(0,0,0,0.15);
        }}
        .image-item img {{
            width: 100%;
            height: 200px;
            object-fit: cover;
            display: block;
        }}
        .image-name {{
            padding: 12px;
            font-size: 14px;
            color: #333;
            text-align: center;
            word-break: break-all;
        }}
        .lightbox {{
            display: none;
            position: fixed;
            top: 0;
            left: 0;
            width: 100%;
            height: 100%;
            background: rgba(0,0,0,0.9);
            z-index: 1000;
            justify-content: center;
            align-items: center;
        }}
        .lightbox.active {{
            display: flex;
        }}
        .lightbox img {{
            max-width: 90%;
            max-height: 90%;
            object-fit: contain;
        }}
        .lightbox-close {{
            position: absolute;
            top: 20px;
            right: 30px;
            color: white;
            font-size: 40px;
            cursor: pointer;
            z-index: 1001;
        }}
        @media (max-width: 768px) {{
            .gallery {{
                grid-template-columns: repeat(auto-fill, minmax(150px, 1fr));
                gap: 10px;
            }}
            .container {{
                padding: 15px;
            }}
        }}
    </style>
</head>
<body>
    <div class="container">
        <h1>📁 {dir_name}</h1>
        <div class="info">共 {len(uploaded_files)} 张图片</div>
        <div class="gallery">
"""
        
        # 添加每张图片
        for file_info in uploaded_files:
            filename = os.path.basename(file_info['local_path'])
            # 直接使用URL，不进行编码（URL中的中文已转换为拼音首字母）
            url = file_info['url']
            
            html_content += f"""
            <div class="image-item" onclick="openLightbox('{url}')">
                <img src="{url}" alt="{filename}" loading="lazy">
                <div class="image-name">{filename}</div>
            </div>
"""
        
        html_content += """
        </div>
    </div>
    
    <div class="lightbox" id="lightbox" onclick="closeLightbox()">
        <span class="lightbox-close">&times;</span>
        <img id="lightbox-img" src="" alt="">
    </div>
    
    <script>
        function openLightbox(url) {
            document.getElementById('lightbox').classList.add('active');
            document.getElementById('lightbox-img').src = url;
        }
        
        function closeLightbox() {
            document.getElementById('lightbox').classList.remove('active');
        }
        
        document.addEventListener('keydown', function(e) {
            if (e.key === 'Escape') {
                closeLightbox();
            }
        });
    </script>
</body>
</html>
"""
        
        # 保存HTML文件
        index_path = os.path.join(directory, 'index.html')
        try:
            with open(index_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            return index_path
        except Exception as e:
            self.log(f"  生成index.html失败: {str(e)}")
            return None
    
    def upload_merged_image_to_oss(self, merged_image_path, directory, root_dir=None):
        """
        上传合并后的大图到OSS
        
        Args:
            merged_image_path: 合并后的图片路径
            directory: 原目录路径
            root_dir: 根目录路径（用于构建完整路径结构）
            
        Returns:
            (成功, 图片URL)
        """
        if not self.oss_uploader:
            self.log("  OSS未配置，跳过上传")
            return False, None
        
        # 构建OSS路径：包含根目录名称和子目录名称，并转换为拼音首字母
        if root_dir:
            root_name = os.path.basename(root_dir)
            dir_name = os.path.basename(directory)
            oss_path = f"{root_name}/{dir_name}/{os.path.basename(merged_image_path)}"
        else:
            dir_name = os.path.basename(directory)
            oss_path = f"{dir_name}/{os.path.basename(merged_image_path)}"
        
        # 将OSS路径中的中文转换为拼音首字母
        pinyin_oss_path = convert_path_to_pinyin(oss_path)
        
        self.log(f"  开始上传合并图片到OSS...")
        success, result = self.oss_uploader.upload_file(merged_image_path, pinyin_oss_path)
        
        if success:
            self.log(f"    ✓ 已上传: {os.path.basename(merged_image_path)}")
            return True, result
        else:
            self.log(f"    ✗ 上传失败: {result}")
            return False, None
    
    def get_directory_prefix(self, directory, root_dir):
        """
        获取目录的完整路径前缀（用于文件命名）
        
        Args:
            directory: 目标目录路径
            root_dir: 根目录路径
            
        Returns:
            目录前缀字符串，例如 "根目录_一级_二级"
        """
        if root_dir and directory.startswith(root_dir):
            root_name = os.path.basename(root_dir)
            rel_path = os.path.relpath(directory, root_dir)
            if rel_path == '.':
                return root_name
            else:
                # 将路径分隔符替换为下划线
                path_parts = [root_name] + rel_path.split(os.sep)
                return '_'.join(path_parts)
        else:
            return os.path.basename(directory)
    
    def copy_to_error_folder(self, directory, root_dir, error_message):
        """
        将处理失败的目录复制到error文件夹，保持原有目录结构
        
        Args:
            directory: 失败的目录路径
            root_dir: 根目录路径
            error_message: 错误信息
        """
        try:
            import shutil
            
            # 创建error文件夹（在根目录同级）
            parent_dir = os.path.dirname(root_dir) if root_dir else os.path.dirname(directory)
            error_base = os.path.join(parent_dir, "error")
            
            # 计算相对路径并创建对应的error子目录
            if root_dir and directory.startswith(root_dir):
                rel_path = os.path.relpath(directory, root_dir)
                error_target = os.path.join(error_base, os.path.basename(root_dir), rel_path)
            else:
                error_target = os.path.join(error_base, os.path.basename(directory))
            
            # 如果目标已存在，先删除
            if os.path.exists(error_target):
                shutil.rmtree(error_target)
            
            # 复制目录（而不是移动）
            shutil.copytree(directory, error_target)
            
            self.log(f"  ✗ 失败：{error_message}")
            self.log(f"     已复制至: {error_target}")
            
        except Exception as e:
            self.log(f"  复制到error文件夹失败: {str(e)}")
    
    def process_directory(self, directory, doc_type, page_size, page_width_mm, page_height_mm, 
                         qr_size_mm, x_mm, y_mm, auto_upload=False, root_dir=None):
        """
        处理单个目录：上传图片、生成二维码和文档（PDF或Word）
        
        Args:
            directory: 目标目录路径
            doc_type: 文档类型（"PDF"或"Word"）
            page_size: PDF页面尺寸（仅PDF使用）
            page_width_mm: Word页面宽度（毫米，仅Word使用）
            page_height_mm: Word页面高度（毫米，仅Word使用）
            qr_size_mm: 二维码大小
            x_mm: 二维码X坐标
            y_mm: 二维码Y坐标
            auto_upload: 是否自动上传
            root_dir: 根目录路径（用于构建OSS路径）
            
        Returns:
            (success: bool, error_message: str) 成功返回(True, None)，失败返回(False, 错误信息)
        """
        dir_name = os.path.basename(directory)
        dir_prefix = self.get_directory_prefix(directory, root_dir)
        
        try:
            # 检查目录中是否有图片
            images = self.get_images_in_directory(directory)
            if not images:
                return True, None  # 无图片不算错误
            
            # 合并图片
            merged_filename = f"{dir_prefix}_merged.jpg"
            merged_path = os.path.join(directory, merged_filename)
            
            if not self.merge_images_horizontally(images, merged_path):
                return False, "图片合并失败"
            
            # 如果启用自动上传，上传合并后的大图到OSS
            oss_url = None
            if auto_upload:
                success, oss_url = self.upload_merged_image_to_oss(merged_path, directory, root_dir)
                if not success:
                    return False, "OSS上传失败"
        
            # 如果没有OSS URL，使用默认格式
            if not oss_url:
                # 构建默认OSS URL
                if self.oss_config.is_valid():
                    endpoint_without_protocol = self.oss_config.endpoint.replace('http://', '').replace('https://', '')
                    
                    # 构建OSS路径
                    if root_dir:
                        root_name = os.path.basename(root_dir)
                        # 计算相对路径
                        if directory.startswith(root_dir):
                            rel_path = os.path.relpath(directory, root_dir)
                            if rel_path == '.':
                                oss_path = f"{root_name}/{merged_filename}"
                            else:
                                oss_path = f"{root_name}/{rel_path}/{merged_filename}"
                        else:
                            oss_path = f"{dir_name}/{merged_filename}"
                    else:
                        oss_path = f"{dir_name}/{merged_filename}"
                    
                    # 将路径中的中文转换为拼音首字母
                    pinyin_path = convert_path_to_pinyin(oss_path)
                    
                    # 构建完整URL
                    if self.oss_config.base_path:
                        pinyin_base = convert_path_to_pinyin(self.oss_config.base_path.strip('/'))
                        oss_url = f"https://{self.oss_config.bucket_name}.{endpoint_without_protocol}/{pinyin_base}/{pinyin_path}"
                    else:
                        oss_url = f"https://{self.oss_config.bucket_name}.{endpoint_without_protocol}/{pinyin_path}"
                else:
                    return False, "OSS未配置"
            
            # 生成二维码
            qr_filename = f"{dir_prefix}_qr.png"
            qr_path = os.path.join(directory, qr_filename)
            
            if not self.generate_qrcode(oss_url, qr_path, qr_size_mm):
                return False, "二维码生成失败"
            
            # 根据文档类型生成PDF或Word
            if doc_type == "PDF":
                doc_filename = f"{dir_prefix}_qr.pdf"
                doc_path = os.path.join(directory, doc_filename)
                
                if not self.create_pdf_with_qrcode(qr_path, doc_path, page_size, qr_size_mm, x_mm, y_mm):
                    return False, "PDF生成失败"
            else:  # Word
                doc_filename = f"{dir_prefix}_qr.docx"
                doc_path = os.path.join(directory, doc_filename)
                
                if not self.create_word_with_qrcode(qr_path, doc_path, page_width_mm, page_height_mm, 
                                                    qr_size_mm, x_mm, y_mm):
                    return False, "Word生成失败"
            
            return True, None
            
        except Exception as e:
            return False, f"处理异常: {str(e)}"
    
    def start_processing(self):
        """开始处理"""
        # 验证输入
        root_dir = self.root_dir_var.get()
        if not root_dir or not os.path.isdir(root_dir):
            messagebox.showerror("错误", "请选择有效的根目录")
            return
        
        # 获取文档类型
        doc_type = self.doc_type_var.get()
        
        # 根据文档类型获取页面尺寸
        if doc_type == "PDF":
            page_size = self.get_page_size()
            if page_size is None:
                return
            page_width_mm = None
            page_height_mm = None
        else:  # Word
            page_size = None
            page_width_mm, page_height_mm = self.get_word_page_size()
            if page_width_mm is None or page_height_mm is None:
                return
        
        try:
            qr_size_mm = float(self.qr_size_var.get())
            x_mm = float(self.qr_x_var.get())
            y_mm = float(self.qr_y_var.get())
        except ValueError:
            messagebox.showerror("错误", "二维码大小和坐标必须是数字")
            return
        
        auto_upload = self.auto_upload_var.get()
        
        if auto_upload and not self.oss_config.is_valid():
            doc_name = "文档" if doc_type == "Word" else "PDF"
            result = messagebox.askyesno("OSS未配置", 
                                        f"你选择了自动上传，但OSS未配置。\n是否继续（仅生成二维码和{doc_name}）？")
            if not result:
                return
            auto_upload = False
        
        # 在新线程中处理，避免阻塞GUI
        thread = threading.Thread(target=self.process_all_directories,
                                 args=(root_dir, doc_type, page_size, page_width_mm, page_height_mm, 
                                       qr_size_mm, x_mm, y_mm, auto_upload))
        thread.daemon = True
        thread.start()
    
    def process_all_directories(self, root_dir, doc_type, page_size, page_width_mm, page_height_mm, 
                                qr_size_mm, x_mm, y_mm, auto_upload):
        """处理所有目录（在后台线程中运行）"""
        try:
            # 禁用按钮
            self.start_button.config(state='disabled')
            self.progress_bar.start()
            self.progress_var.set("正在处理...")
            
            self.log("=" * 60)
            self.log("开始处理...")
            self.log(f"根目录: {root_dir}")
            self.log(f"目录类型: {self.dir_type_var.get()}")
            self.log(f"文档类型: {doc_type}")
            if doc_type == "PDF":
                self.log(f"页面尺寸: {self.page_size_var.get()}")
            else:
                self.log(f"页面尺寸: {self.word_size_var.get()} ({page_width_mm}mm x {page_height_mm}mm)")
            self.log(f"二维码大小: {qr_size_mm}mm")
            self.log(f"二维码位置: ({x_mm}mm, {y_mm}mm)")
            self.log(f"自动上传: {'是' if auto_upload else '否'}")
            self.log("=" * 60)
            
            # 获取目标目录
            dir_type = self.dir_type_var.get()
            target_dirs = self.get_target_directories(root_dir, dir_type)
            
            self.log(f"找到 {len(target_dirs)} 个目标目录")
            self.log("")
            
            # 处理每个目录
            success_count = 0
            fail_count = 0
            skip_count = 0
            
            for target_dir in target_dirs:
                dir_name = os.path.basename(target_dir)
                
                # 检查是否有图片
                images = self.get_images_in_directory(target_dir)
                if not images:
                    self.log(f"跳过 {dir_name} (无图片)")
                    skip_count += 1
                    continue
                
                # 处理目录
                success, error_msg = self.process_directory(target_dir, doc_type, page_size, page_width_mm, 
                                                           page_height_mm, qr_size_mm, x_mm, y_mm, 
                                                           auto_upload, root_dir)
                
                if success:
                    self.log(f"✓ 成功: {dir_name}")
                    success_count += 1
                else:
                    self.log(f"✗ 失败: {dir_name} - {error_msg}")
                    fail_count += 1
                    # 复制到error文件夹
                    self.copy_to_error_folder(target_dir, root_dir, error_msg)
            
            self.log("")
            self.log("=" * 60)
            self.log(f"处理完成！")
            self.log(f"  成功: {success_count} 个")
            self.log(f"  失败: {fail_count} 个")
            self.log(f"  跳过: {skip_count} 个")
            self.log("=" * 60)
            
            self.progress_var.set("处理完成")
            messagebox.showinfo("完成", f"处理完成！\n成功: {success_count}\n失败: {fail_count}\n跳过: {skip_count}")
            
        except Exception as e:
            self.log(f"处理过程中出错: {str(e)}")
            messagebox.showerror("错误", f"处理过程中出错: {str(e)}")
        finally:
            # 恢复按钮
            self.start_button.config(state='normal')
            self.progress_bar.stop()


def main():
    """主函数"""
    root = tk.Tk()
    app = DocumentProcessorApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
